"""Extraction et analyse des images dans les documents."""
import logging
from pathlib import Path
from typing import List, Dict
import re

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

logger = logging.getLogger(__name__)


class ImageExtractor:
    """Extrait les métadonnées des images des documents."""

    def extract_images_metadata(self, file_path: Path, text: str) -> Dict:
        """Extrait les métadonnées des images d'un document.

        Args:
            file_path: Chemin vers le fichier
            text: Texte extrait du document (pour corrélation)

        Returns:
            Dictionnaire avec métadonnées des images
        """
        file_ext = file_path.suffix.lower()

        if file_ext == ".docx":
            return self._extract_from_docx(file_path, text)
        elif file_ext == ".pdf":
            return self._extract_from_pdf(file_path, text)
        else:
            return {"image_count": 0, "images": [], "error": "Format non supporté"}

    def _extract_from_docx(self, file_path: Path, text: str) -> Dict:
        """Extrait les images d'un fichier DOCX.

        Args:
            file_path: Chemin vers le fichier DOCX
            text: Texte du document

        Returns:
            Métadonnées des images
        """
        try:
            doc = Document(str(file_path))
            images = []
            image_count = 0

            # Parcourir les relations du document pour trouver les images
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_count += 1

            # Essayer de localiser les images dans le document
            # (approximatif car python-docx ne donne pas facilement la position exacte)
            paragraph_count = len(doc.paragraphs)
            approx_position = 0

            for i, paragraph in enumerate(doc.paragraphs):
                # Vérifier si le paragraphe contient une image (via les runs)
                for run in paragraph.runs:
                    if run._element.xpath('.//a:blip'):
                        # Image trouvée dans ce paragraphe
                        # Capturer le contexte textuel (paragraphes autour)
                        context_before = []
                        context_after = []

                        # 3 paragraphes avant
                        for j in range(max(0, i-3), i):
                            if doc.paragraphs[j].text.strip():
                                context_before.append(doc.paragraphs[j].text.strip())

                        # 3 paragraphes après
                        for j in range(i+1, min(paragraph_count, i+4)):
                            if doc.paragraphs[j].text.strip():
                                context_after.append(doc.paragraphs[j].text.strip())

                        context = " ".join(context_before + context_after)

                        images.append({
                            "position": i,
                            "approx_location": f"Paragraphe {i}/{paragraph_count}",
                            "context": context[:300]  # Premier 300 caractères de contexte
                        })

            return {
                "image_count": image_count,
                "images": images,
                "source": "docx"
            }

        except Exception as e:
            logger.error(f"Erreur extraction images DOCX: {e}")
            return {"image_count": 0, "images": [], "error": str(e)}

    def _extract_from_pdf(self, file_path: Path, text: str) -> Dict:
        """Extrait les images d'un fichier PDF.

        Args:
            file_path: Chemin vers le fichier PDF
            text: Texte du document

        Returns:
            Métadonnées des images
        """
        if PdfReader is None:
            return {"image_count": 0, "images": [], "error": "PyPDF2 non installé"}

        try:
            reader = PdfReader(str(file_path))
            image_count = 0
            images = []

            for page_num, page in enumerate(reader.pages, 1):
                # Compter les images sur cette page
                if '/XObject' in page['/Resources']:
                    xobjects = page['/Resources']['/XObject'].get_object()
                    for obj in xobjects:
                        if xobjects[obj]['/Subtype'] == '/Image':
                            image_count += 1

                            # Essayer d'extraire le texte de cette page pour contexte
                            try:
                                page_text = page.extract_text()
                                context = page_text[:300] if page_text else "Contexte non disponible"
                            except:
                                context = "Contexte non disponible"

                            images.append({
                                "page": page_num,
                                "approx_location": f"Page {page_num}/{len(reader.pages)}",
                                "context": context
                            })

            return {
                "image_count": image_count,
                "images": images,
                "source": "pdf"
            }

        except Exception as e:
            logger.error(f"Erreur extraction images PDF: {e}")
            return {"image_count": 0, "images": [], "error": str(e)}

    def detect_illustration_zones(self, text: str) -> List[Dict]:
        """Détecte les zones de texte qui devraient contenir des illustrations.

        Args:
            text: Texte complet du document

        Returns:
            Liste des zones détectées
        """
        zones = []

        # Patterns indiquant une illustration probable
        patterns = [
            # Listes de méthodes/étapes
            r"(de la manière suivante|comme suit|ci-dessous|voir (schéma|figure|photo))[\s\S]{0,50}[-•\d\.]",
            # Sections méthodologiques
            r"(protocole|méthode|procédure|installation)[\s\S]{0,100}[-•\d\.][\s\S]{0,100}[-•\d\.]",
            # Mentions explicites
            r"(figure|schéma|illustration|photo|diagramme)\s+\d+",
            # Découpe/installation
            r"(découp|install|positionn|mesur)[\s\S]{0,50}(suivant|ci-dessous|comme suit)"
        ]

        text_lower = text.lower()

        for pattern in patterns:
            matches = re.finditer(pattern, text_lower, re.IGNORECASE)
            for match in matches:
                start = match.start()
                end = match.end()

                # Capturer contexte étendu (200 caractères avant et après)
                context_start = max(0, start - 200)
                context_end = min(len(text), end + 200)
                context = text[context_start:context_end]

                # Vérifier si déjà capturé (éviter doublons)
                if not any(abs(z["position"] - start) < 100 for z in zones):
                    zones.append({
                        "position": start,
                        "matched_pattern": pattern,
                        "context": context,
                        "keywords": self._extract_keywords(context)
                    })

        return zones

    def _extract_keywords(self, context: str) -> List[str]:
        """Extrait les mots-clés techniques du contexte.

        Args:
            context: Texte de contexte

        Returns:
            Liste de mots-clés
        """
        technical_words = [
            "carotte", "découpe", "compression", "carbonatation", "chlorure",
            "essai", "mesure", "protocole", "schéma", "méthode",
            "installation", "instrumentation", "capteur", "accéléromètre",
            "géoradar", "ferroscan", "corrosimètre", "auscultation"
        ]

        context_lower = context.lower()
        found_keywords = [
            word for word in technical_words
            if word in context_lower
        ]

        return found_keywords

    def correlate_zones_and_images(
        self, zones: List[Dict], images_metadata: Dict, text: str
    ) -> List[Dict]:
        """Corrèle les zones textuelles avec les images détectées.

        Args:
            zones: Zones textuelles détectées
            images_metadata: Métadonnées des images extraites
            text: Texte complet du document

        Returns:
            Liste des illustrations exceptionnelles détectées
        """
        illustrations = []
        images = images_metadata.get("images", [])
        image_count = images_metadata.get("image_count", 0)

        if image_count == 0:
            # Pas d'images, mais on peut quand même signaler les zones probables
            for zone in zones:
                if len(zone["keywords"]) >= 2:  # Au moins 2 mots-clés techniques
                    illustrations.append({
                        "type": "illustration probable (image non détectée)",
                        "description": self._generate_description(zone),
                        "context": zone["context"][:200],
                        "detection_method": "analyse textuelle",
                        "confidence": "low",
                        "keywords": zone["keywords"]
                    })
        else:
            # Corréler zones et images
            for zone in zones:
                # Chercher l'image la plus proche (approximatif)
                closest_image = self._find_closest_image(zone, images, text)

                if closest_image:
                    illustrations.append({
                        "type": "illustration méthodologique (détectée)",
                        "description": self._generate_description(zone),
                        "context": zone["context"][:200],
                        "location": closest_image.get("approx_location", "Position inconnue"),
                        "detection_method": "corrélation texte + image",
                        "confidence": "high" if len(zone["keywords"]) >= 3 else "medium",
                        "keywords": zone["keywords"]
                    })
                elif len(zone["keywords"]) >= 3:
                    # Zone très technique mais pas d'image proche trouvée
                    illustrations.append({
                        "type": "illustration probable (non corrélée)",
                        "description": self._generate_description(zone),
                        "context": zone["context"][:200],
                        "detection_method": "analyse textuelle forte",
                        "confidence": "medium",
                        "keywords": zone["keywords"]
                    })

        return illustrations

    def _find_closest_image(self, zone: Dict, images: List[Dict], text: str) -> Dict:
        """Trouve l'image la plus proche d'une zone textuelle.

        Args:
            zone: Zone textuelle
            images: Liste des images
            text: Texte complet

        Returns:
            Image la plus proche ou None
        """
        if not images:
            return None

        zone_position = zone["position"]

        # Pour DOCX : comparer les positions de paragraphes
        # Pour PDF : comparer le contexte textuel
        best_match = None
        best_score = 0

        for image in images:
            score = 0

            # Vérifier si le contexte de l'image contient des mots-clés de la zone
            image_context = image.get("context", "").lower()
            zone_context = zone["context"].lower()

            # Compter mots communs
            zone_words = set(zone_context.split())
            image_words = set(image_context.split())
            common_words = zone_words & image_words

            score = len(common_words)

            # Bonus si keywords présents dans contexte image
            for keyword in zone.get("keywords", []):
                if keyword in image_context:
                    score += 5

            if score > best_score:
                best_score = score
                best_match = image

        # Retourner seulement si score significatif
        return best_match if best_score > 3 else None

    def _generate_description(self, zone: Dict) -> str:
        """Génère une description basée sur la zone détectée.

        Args:
            zone: Zone textuelle

        Returns:
            Description générée
        """
        keywords = zone.get("keywords", [])
        context = zone.get("context", "")

        if not keywords:
            return "Illustration méthodologique non spécifiée"

        # Essayer de construire une description basée sur les mots-clés
        if "carotte" in keywords and "découpe" in keywords:
            tests = [k for k in keywords if k in ["compression", "carbonatation", "chlorure"]]
            if tests:
                return f"Protocole de découpe de carotte pour essais ({', '.join(tests)})"
            return "Protocole de découpe de carotte"

        if "installation" in keywords or "instrumentation" in keywords:
            devices = [k for k in keywords if k in ["capteur", "accéléromètre", "géoradar"]]
            if devices:
                return f"Installation/Configuration {', '.join(devices)}"
            return "Schéma d'installation/instrumentation"

        # Description générique basée sur les mots-clés
        return f"Illustration méthodologique ({', '.join(keywords[:3])})"
