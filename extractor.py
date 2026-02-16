"""Extraction de texte depuis les fichiers Word (.docx) et PDF."""
import logging
from pathlib import Path
from typing import Optional, Tuple

from docx import Document
from PyPDF2 import PdfReader

logger = logging.getLogger(__name__)


class TextExtractor:
    """Extrait le texte des fichiers Word et PDF."""

    @staticmethod
    def extract_from_docx(file_path: Path) -> Optional[str]:
        """Extrait le texte d'un fichier .docx.

        Args:
            file_path: Chemin vers le fichier .docx

        Returns:
            Le texte extrait ou None en cas d'erreur
        """
        try:
            doc = Document(file_path)
            text_parts = []

            # Extraire le texte des paragraphes
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extraire le texte des tableaux
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            return "\n".join(text_parts)

        except Exception as e:
            logger.error(f"Erreur lors de l'extraction de {file_path}: {e}")
            return None

    @staticmethod
    def extract_from_pdf(file_path: Path) -> Optional[str]:
        """Extrait le texte d'un fichier PDF.

        Args:
            file_path: Chemin vers le fichier PDF

        Returns:
            Le texte extrait ou None en cas d'erreur
        """
        try:
            reader = PdfReader(file_path)
            text_parts = []

            for page in reader.pages:
                text = page.extract_text()
                if text.strip():
                    text_parts.append(text)

            return "\n".join(text_parts)

        except Exception as e:
            logger.error(f"Erreur lors de l'extraction de {file_path}: {e}")
            return None

    @staticmethod
    def get_page_count_docx(file_path: Path) -> Optional[int]:
        """Obtient le nombre de pages d'un fichier .docx.

        Args:
            file_path: Chemin vers le fichier .docx

        Returns:
            Le nombre de pages ou None en cas d'erreur
        """
        try:
            doc = Document(file_path)
            # Essayer d'obtenir le nombre de pages depuis les propriétés du document
            if hasattr(doc.core_properties, 'pages') and doc.core_properties.pages:
                return doc.core_properties.pages

            # Sinon, estimer basé sur le nombre de paragraphes (approximatif)
            # En moyenne, ~40-50 lignes par page, et ~2-3 paragraphes par ligne
            # Donc environ 15-20 paragraphes par page
            total_paragraphs = len([p for p in doc.paragraphs if p.text.strip()])
            estimated_pages = max(1, total_paragraphs // 15)
            return estimated_pages

        except Exception as e:
            logger.error(f"Erreur lors de l'obtention du nombre de pages de {file_path}: {e}")
            return None

    @staticmethod
    def get_page_count_pdf(file_path: Path) -> Optional[int]:
        """Obtient le nombre de pages d'un fichier PDF.

        Args:
            file_path: Chemin vers le fichier PDF

        Returns:
            Le nombre de pages ou None en cas d'erreur
        """
        try:
            reader = PdfReader(file_path)
            return len(reader.pages)
        except Exception as e:
            logger.error(f"Erreur lors de l'obtention du nombre de pages de {file_path}: {e}")
            return None

    @classmethod
    def get_page_count(cls, file_path: Path) -> Optional[int]:
        """Obtient le nombre de pages d'un fichier selon son extension.

        Args:
            file_path: Chemin vers le fichier

        Returns:
            Le nombre de pages ou None en cas d'erreur
        """
        suffix = file_path.suffix.lower()

        if suffix in [".docx", ".doc", ".docm"]:
            return cls.get_page_count_docx(file_path)
        elif suffix == ".pdf":
            return cls.get_page_count_pdf(file_path)
        else:
            logger.warning(f"Extension non supportée: {suffix}")
            return None

    @classmethod
    def extract(cls, file_path: Path) -> Optional[str]:
        """Extrait le texte d'un fichier selon son extension.

        Args:
            file_path: Chemin vers le fichier

        Returns:
            Le texte extrait ou None en cas d'erreur
        """
        suffix = file_path.suffix.lower()

        if suffix in [".docx", ".doc", ".docm"]:
            # python-docx peut lire .docx, .doc (Word 97-2003) et .docm (avec macros)
            return cls.extract_from_docx(file_path)
        elif suffix == ".pdf":
            return cls.extract_from_pdf(file_path)
        else:
            logger.warning(f"Extension non supportée: {suffix}")
            return None

    @classmethod
    def extract_with_metadata(cls, file_path: Path) -> Tuple[Optional[str], Optional[int]]:
        """Extrait le texte et le nombre de pages d'un fichier.

        Args:
            file_path: Chemin vers le fichier

        Returns:
            Tuple (texte, nombre_de_pages) ou (None, None) en cas d'erreur
        """
        text = cls.extract(file_path)
        page_count = cls.get_page_count(file_path)
        return text, page_count


def main():
    """Test du module d'extraction."""
    import sys

    if len(sys.argv) < 2:
        print("Usage: python extractor.py <chemin_fichier>")
        sys.exit(1)

    file_path = Path(sys.argv[1])

    if not file_path.exists():
        print(f"Fichier introuvable: {file_path}")
        sys.exit(1)

    extractor = TextExtractor()
    text = extractor.extract(file_path)

    if text:
        print(f"=== Texte extrait de {file_path.name} ===")
        print(text[:500] + "..." if len(text) > 500 else text)
        print(f"\nLongueur totale: {len(text)} caractères")
    else:
        print("Échec de l'extraction")


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    main()
